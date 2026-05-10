"""
Genera el documento de referencia técnica del dashboard "Engordando Holando".

Salida: "Parametros y formulas de dashboard de engorde holando.docx" en la raíz
del proyecto. Re-ejecutable; sobreescribe el archivo si existe.
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_TITLE = "Parámetros y fórmulas — Dashboard Engordando Holando"
OUT_PATH = Path(__file__).parent / "Parametros y formulas de dashboard de engorde holando.docx"

C_PRIMARY = RGBColor(0x0C, 0x1A, 0x2E)
C_ACCENT  = RGBColor(0x15, 0x65, 0xC0)
C_MUTED   = RGBColor(0x5D, 0x7A, 0x95)
C_SUBTLE  = RGBColor(0x94, 0xA3, 0xB8)


# ── Helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = C_PRIMARY


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = C_ACCENT


def add_h3(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = C_PRIMARY


def add_p(doc: Document, text: str, *, italic: bool = False, size: float = 10.5) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10.5)
        # add the text:
        r = p.runs[0] if p.runs else p.add_run("")
        if not p.runs or not p.runs[0].text:
            r.text = it
        else:
            r.text = it
        r.font.size = Pt(10.5)


def add_formula(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.font.color.rgb = C_PRIMARY


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
               col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False

    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, "1565C0")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(str(val))
            run.font.size = Pt(9.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


def add_pagebreak(doc: Document) -> None:
    doc.add_page_break()


# ── Construcción del documento ─────────────────────────────────────────────

doc = Document()

# Margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

# ── Portada ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
r = p.add_run("Engordando Holando")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = C_PRIMARY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Parámetros y fórmulas del dashboard de engorde Holando")
r.font.size = Pt(14)
r.font.color.rgb = C_ACCENT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run("Documentación técnica · referencia de fórmulas · guía de uso")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = C_MUTED

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(220)
r = p.add_run("by Sofía Stieglitz · mayo 2026")
r.font.size = Pt(10)
r.font.color.rgb = C_SUBTLE

add_pagebreak(doc)


# ── 1. INTRODUCCIÓN ────────────────────────────────────────────────────────

add_h1(doc, "1. Introducción")

add_h2(doc, "1.1 Objetivo del dashboard")
add_p(doc,
      "El dashboard simula la cadena productiva y económica del engorde de "
      "terneros macho Holando desde el destete hasta la terminación. Permite "
      "evaluar márgenes, costos, sensibilidad y robustez del modelo bajo "
      "distintas combinaciones de etapas productivas, parámetros bioeconómicos "
      "y escenarios de precios.")

add_h2(doc, "1.2 Etapas productivas")
add_p(doc, "El ciclo se descompone en tres etapas:")
add_table(doc,
    headers=["Etapa", "Rango (kg)", "Días típicos", "Caracterización"],
    rows=[
        ["Cría (A)",       "45 → 95",   "60",  "Recibo del ternero, iniciador, alta sanidad."],
        ["Recría (B)",     "95 → 220",  "200", "Pastura + suplemento, GDP medio, transición."],
        ["Engorde (C)",    "220 → 430", "180", "Feedlot, ración alta energía, terminación."],
    ],
    col_widths=[3.0, 2.5, 2.5, 8.0],
)

add_h2(doc, "1.3 Lógica modular y encadenamiento")
add_p(doc,
      "El dashboard puede correr en modo integrado (Cría → Recría → Engorde) "
      "o modular: cualquier slice contiguo de etapas activas — solo Cría, "
      "solo Recría, solo Engorde, Cría+Recría o Recría+Engorde. La "
      "combinación Cría+Engorde sin Recría no es válida y se fuerza a "
      "contiguidad agregando Recría automáticamente.")
add_bullets(doc, [
    "El kg de entrada de la primera etapa activa es input editable.",
    "Las etapas encadenadas heredan kg_in del kg_out de la etapa previa activa.",
    "Las cabezas en cascada respetan únicamente el slice activo: la primera "
    "etapa parte de n_terneros; cada siguiente recibe los sobrevivientes "
    "de la anterior.",
    "Las etapas inactivas se muestran atenuadas (opacity 0.42 + badge "
    "INACTIVA) pero no participan de los cálculos agregados.",
])

add_h2(doc, "1.4 Enfoque bioeconómico")
add_p(doc,
      "La alimentación se modela siempre de forma bioeconómica pura: el costo "
      "se deriva de los kilos de carne producidos (Δkg × CA × precio "
      "ponderado de la ración). No existen entradas USD/cab/día manuales — "
      "la fuente de verdad es la tabla de feed editable por etapa "
      "(ingredientes con % de inclusión y precio USD/kg MS).")
add_p(doc,
      "Mortandad se contabiliza como ingreso perdido: las cabezas que mueren "
      "no generan venta, pero sí incurren costo durante el período (modelo "
      "asimétrico cab_in → cab_vend).", italic=True)

add_pagebreak(doc)


# ── 2. SLIDE POR SLIDE ─────────────────────────────────────────────────────

add_h1(doc, "2. Slide por slide")

# ── 2.1 Parámetros ─────────────────────────────────────────────────────────
add_h2(doc, "2.1 Parámetros")
add_h3(doc, "Objetivo")
add_p(doc,
      "Configurar todos los inputs del modelo: biológicos, sanidad, "
      "alimentación, operación, estructura, comercialización y financieros. "
      "Es el panel de control único; las demás slides leen exclusivamente "
      "de session_state.")
add_h3(doc, "Componentes principales")
add_bullets(doc, [
    "Toggles globales de etapas activas (Cría / Recría / Engorde).",
    "Bloque ‘Comunes’: cantidad de terneros, peso entrada, precio compra, "
    "tipo de cambio, tasa de interés.",
    "Bloque por etapa (A/B/C): kg entrada/salida, días, GDP, CA, mortandad, "
    "sanidad, MO+combustible+servicios mensuales, comercialización.",
    "Tabla editable de feed por etapa: ingredientes con % de inclusión y "
    "USD/kg MS. Es la única fuente del costo de alimentación.",
    "Bloque ‘Estructura’: valor total de infra + asignación % por etapa, "
    "años de amortización y mantenimiento anual.",
])
add_h3(doc, "Lógica de cálculo")
add_p(doc,
      "Los toggles disparan una validación de contiguidad: si el usuario deja "
      "{Cría=ON, Recría=OFF, Engorde=ON} el sistema fuerza Recría a ON. "
      "El kg de entrada editable se libera sólo en la etapa que sea primera "
      "activa; en las encadenadas se deriva automáticamente del kg de salida "
      "de la etapa previa.")

# ── 2.2 Modelo Productivo ──────────────────────────────────────────────────
add_h2(doc, "2.2 Modelo Productivo")
add_h3(doc, "Objetivo")
add_p(doc,
      "Visualizar el ciclo biológico: kg ganados, días, GDP, conversión y "
      "mortandad por etapa. Es la representación productiva sin componente "
      "económico.")
add_h3(doc, "Indicadores")
add_table(doc,
    headers=["Indicador", "Unidad", "Cálculo"],
    rows=[
        ["kg entrada / salida",   "kg",       "Input (1ª activa) o herencia (encadenada)"],
        ["Δkg ganado",            "kg",       "kg_out − kg_in"],
        ["Días",                  "días",     "Input"],
        ["GDP",                   "kg/día",   "(kg_out − kg_in) / días"],
        ["CA",                    "kg MS/kg", "Input"],
        ["Mortandad",             "%",        "Input"],
        ["Cabezas que entran",    "cab",      "Cascada del slice activo"],
        ["Cabezas que egresan",   "cab",      "cab_in × (1 − mort/100), entero"],
    ],
    col_widths=[5.0, 3.0, 8.0],
)

# ── 2.3 Costos ─────────────────────────────────────────────────────────────
add_h2(doc, "2.3 Costos")
add_h3(doc, "Objetivo")
add_p(doc,
      "Desglosar los costos por etapa en 8 categorías. Usa la suma costos "
      "directos + financiero + mortandad como ‘oportunidad perdida’ para "
      "exponer el verdadero costo económico.")
add_h3(doc, "Categorías (USD por cabeza)")
add_table(doc,
    headers=["Bucket", "Fórmula", "Comentario"],
    rows=[
        ["Compra",          "pc_global × kg_in", "Sólo en la 1ª etapa activa con compra externa."],
        ["Alimentación",    "(kg_out − kg_in) × CA × precio_pond_MS", "Bioeconómica."],
        ["Sanidad",         "K.*_SANIDAD",       "USD/cab fijo por etapa."],
        ["Operación",       "(MO + Combust. + Serv.)/30 × días / cabezas", "MO/comb/serv son USD/mes absolutos."],
        ["Estructura",      "(infra × asig%/100 / años + mant_año) × días/365 / cabezas", "Asignación parcial."],
        ["Comercialización","(com%/100) × pv × kg_out + flete_entrada + flete_salida", "Comisión proporcional."],
        ["Financiero",      "(compra+alim+san+op+estr+com) × tasa%/100 × días/365", "Sobre capital total."],
        ["Mortandad",       "(mort%/100) × kg_out × pv", "Ingreso perdido (oportunidad)."],
    ],
    col_widths=[3.5, 6.5, 6.0],
)
add_h3(doc, "Indicadores resumidos")
add_bullets(doc, [
    "USD totales del sistema (suma de etapas activas).",
    "USD por cabeza ingresada por etapa.",
    "USD por kg producido = total_usd / ((kg_out − kg_in) × cab_vend).",
    "% de cada categoría sobre el costo total.",
])

# ── 2.4 Ingresos ───────────────────────────────────────────────────────────
add_h2(doc, "2.4 Ingresos")
add_h3(doc, "Objetivo")
add_p(doc,
      "Cuantificar el ingreso por venta efectiva en cada etapa. Sólo las "
      "cabezas que sobreviven al período generan ingreso real.")
add_h3(doc, "Fórmulas")
add_formula(doc, "ingreso_cab    = kg_out × precio_venta")
add_formula(doc, "cab_vend       = round( cab_in × (1 − mort/100) )")
add_formula(doc, "ingreso_total  = ingreso_cab × cab_vend")
add_formula(doc, "kg_vendidos    = cab_vend × kg_out")
add_h3(doc, "Indicadores")
add_bullets(doc, [
    "Ingreso por cabeza (USD/cab).",
    "Ingreso total por etapa y agregado (USD).",
    "Kilos vendidos (kg).",
    "Cabezas vendidas (cab).",
])

# ── 2.5 Margen Bruto ───────────────────────────────────────────────────────
add_h2(doc, "2.5 Margen Bruto")
add_h3(doc, "Objetivo")
add_p(doc,
      "Integrar costos e ingresos para obtener el margen económico de cada "
      "etapa y del sistema completo. Las cifras incluyen costo financiero "
      "sobre el capital inmovilizado durante el ciclo.")
add_h3(doc, "Buckets de costo y modelo de mortandad")
add_p(doc,
      "El costo se compone de los mismos seis directos que en la slide Costos "
      "más el financiero. La mortandad NO se suma al costo de la etapa: ya "
      "está implícita en la diferencia entre cab_in (incurre el costo) y "
      "cab_vend (genera el ingreso). Esta asimetría evita la doble "
      "contabilización.")
add_h3(doc, "Fórmulas")
add_formula(doc, "capital     = compra + alim + sanidad + op + estr + com")
add_formula(doc, "financiero  = capital × tasa%/100 × días/365")
add_formula(doc, "costo_cab   = capital + financiero")
add_formula(doc, "costo_total = costo_cab × cab_in")
add_formula(doc, "ingreso_total = ingreso_cab × cab_vend       (= kg_out × pv × cab_vend)")
add_formula(doc, "margen_bruto_total = ingreso_total − costo_total")
add_formula(doc, "margen_bruto_cab   = margen_bruto_total / cab_in")
add_formula(doc, "margen_kg          = margen_bruto_total / kg_prod_total")
add_formula(doc, "kg_prod_total      = (kg_out − kg_in) × cab_vend")
add_formula(doc, "USD/cab/día        = margen_bruto_cab / días")
add_h3(doc, "Indicadores derivados")
add_bullets(doc, [
    "Retorno incremental: USD/cab vs etapa previa activa (cadena).",
    "USD/cab/día: ratio de eficiencia económica diaria.",
    "Margen/kg producido: margen económico por kilo de carne efectiva.",
    "Capital comprometido: costo_total agregado.",
])

# ── 2.6 Sensibilidad y Riesgo ──────────────────────────────────────────────
add_h2(doc, "2.6 Sensibilidad y Riesgo")
add_h3(doc, "Objetivo")
add_p(doc,
      "Cuantificar la fragilidad económica del modelo frente a variaciones "
      "de precios, productividad y riesgo operativo. Combina análisis "
      "tornado, simulación interactiva y un score de robustez por etapa "
      "activa.")
add_h3(doc, "Estructura de la slide")
add_bullets(doc, [
    "Tornado por etapa con 6 variables y rango ±20% (excepto GDP ±15% y "
    "mortandad ±5pp).",
    "Simulación interactiva: el usuario manipula sliders y ve la respuesta.",
    "Robustez por etapa activa: score 0–100 (mayor = más robusto).",
    "Resumen comparativo de breakeven y headroom.",
    "Alertas focalizadas en sensibilidades altas y robustez baja.",
])
add_h3(doc, "Variables del tornado")
add_table(doc,
    headers=["Variable", "Δ aplicado", "Descripción"],
    rows=[
        ["Precio compra",  "±20 %",  "Multiplica el costo de compra de la 1ª etapa activa."],
        ["Precio venta",   "±20 %",  "Multiplica el ingreso y la base de la comisión."],
        ["Precio maíz",    "±20 %",  "Multiplica el componente de alimentación."],
        ["GDP",            "±15 %",  "Escala los kg ganados (y proporcionalmente la alim)."],
        ["Mortandad",      "±5 pp",  "Suma puntos porcentuales sobre la mortandad base."],
        ["Flete",          "±20 %",  "Multiplica fletes de entrada y salida."],
    ],
    col_widths=[3.5, 2.5, 10.0],
)
add_h3(doc, "Score de robustez")
add_p(doc,
      "Robustez = 100 − composite, donde composite es el promedio ponderado "
      "de cinco factores que penalizan exposición al riesgo:")
add_table(doc,
    headers=["Factor", "Cálculo", "Cap"],
    rows=[
        ["Sensibilidad maíz",  "swing(maíz) / |baseline| × 50",   "100"],
        ["Volatilidad",        "Σ swings / |baseline| × 30",       "100"],
        ["Duración",           "días / 7.30",                      "100"],
        ["Capital relativo",   "costo_total / max(activos) × 100", "100"],
        ["Mortandad",          "mort% × 10",                       "100"],
    ],
    col_widths=[4.5, 7.0, 1.5],
)

# ── 2.7 Reportes ───────────────────────────────────────────────────────────
add_h2(doc, "2.7 Reportes")
add_h3(doc, "Objetivo")
add_p(doc,
      "Generar un one-pager ejecutivo con formato A4 — header con "
      "empresa/responsable/fecha, resumen ejecutivo del sistema integrado, "
      "y mini-bloques por slide con bullets y un mini gráfico inline.")
add_h3(doc, "Estructura")
add_bullets(doc, [
    "Toolbar superior editable: empresa, responsable, fecha, logo on/off.",
    "Header A4: logo a izquierda + datos a derecha.",
    "Hero ‘Sistema productivo integrado’ con margen total como headline.",
    "4 KPI cards: Margen sistema, USD/cab/día, Riesgo (promedio), Robustez "
    "(promedio).",
    "Primera fila de mini-bloques: Modelo Productivo, Costos, Ingresos.",
    "Segunda fila de mini-bloques: Margen Bruto, Sensibilidad y Riesgo.",
])

add_pagebreak(doc)


# ── 3. PARÁMETROS ──────────────────────────────────────────────────────────

add_h1(doc, "3. Parámetros")

add_p(doc,
      "Listado de todos los parámetros editables. Todos viven en "
      "session_state y son la única fuente de verdad para las páginas "
      "de cálculo.")

add_h2(doc, "3.1 Comunes")
add_table(doc,
    headers=["Parámetro", "Unidad", "Default", "Tipo", "Descripción"],
    rows=[
        ["n_terneros",     "cab",   "800",     "input",    "Cabezas que ingresan al ciclo."],
        ["peso_inicial",   "kg",    "95",      "input",    "Peso al destete (= kg salida Cría)."],
        ["precio_compra",  "USD/kg","1.20",    "input",    "Precio del ternero ex-tambo."],
        ["tipo_cambio",    "ARS/USD","1100",   "input",    "Referencia mayo 2026."],
        ["tasa_interes",   "% anual","8.0",    "input",    "Costo financiero del capital de trabajo."],
    ],
    col_widths=[3.5, 2.0, 2.0, 1.8, 6.7],
)

add_h2(doc, "3.2 Cría (A) — 45 → 95 kg, 60 días")
add_table(doc,
    headers=["Parámetro", "Unidad", "Default", "Tipo", "Descripción"],
    rows=[
        ["a_kg_entrada",   "kg",       "45",     "input",    "Peso al ingreso a Cría (1ª activa)."],
        ["d_dias",         "días",     "60",     "input",    "Duración del período."],
        ["a_gdp",          "kg/día",   "0.833",  "input",    "Ganancia diaria promedio."],
        ["a_ca",           "kg MS/kg", "4.0",    "input",    "Conversión alimenticia."],
        ["d_mortalidad",   "%",        "4.0",    "input",    "Mortandad esperada de recibo."],
        ["d_sanidad",      "USD/cab",  "22",     "input",    "Vacunas + antiparasitarios + neumonías."],
        ["d_mo_mes",       "USD/mes",  "0",      "input",    "Mano de obra mensual."],
        ["a_combustible",  "USD/mes",  "0",      "input",    "Combustible mensual."],
        ["a_servicios",    "USD/mes",  "0",      "input",    "Servicios mensuales."],
        ["a_asig_pct",     "%",        "0",      "input",    "% de la infra global asignado a Cría."],
        ["a_amort_anos",   "años",     "10",     "input",    "Vida útil de la infra de Cría."],
        ["a_mantenimiento","USD/año",  "0",      "input",    "Mantenimiento anual."],
        ["d_precio_venta", "USD/kg",   "1.40",   "input",    "Precio si la etapa es punto de salida."],
        ["a_fe",           "USD/cab",  "0",      "input",    "Flete de entrada."],
        ["d_flete",        "USD/cab",  "4",      "input",    "Flete de salida."],
        ["a_comision_pct", "%",        "2.0",    "input",    "Comisión sobre ingreso bruto."],
        ["d_otros",        "USD/cab",  "5",      "input",    "Otros gastos varios."],
    ],
    col_widths=[3.7, 2.0, 1.8, 1.6, 7.0],
)

add_h2(doc, "3.3 Recría (B) — 95 → 220 kg, 200 días")
add_table(doc,
    headers=["Parámetro", "Unidad", "Default", "Tipo", "Descripción"],
    rows=[
        ["b_kg_entrada",   "kg",       "95",     "input/herencia", "Editable cuando es 1ª activa; sino hereda kg_out de Cría."],
        ["b_dias",         "días",     "200",    "input",          "Duración del período."],
        ["r_peso_salida",  "kg",       "220",    "input",          "Kg al final de Recría."],
        ["r_gdp",          "kg/día",   "0.625",  "input",          "GDP en pastoreo + suplemento."],
        ["r_ca",           "kg MS/kg", "9.5",    "input",          "Conversión alimenticia."],
        ["r_mortalidad",   "%",        "2.0",    "input",          "Mortandad esperada."],
        ["r_sanidad",      "USD/cab",  "25",     "input",          "Sanidad por cabeza."],
        ["r_mo_mes",       "USD/mes",  "0",      "input",          "MO mensual."],
        ["b_combustible",  "USD/mes",  "0",      "input",          "Combustible."],
        ["b_servicios",    "USD/mes",  "0",      "input",          "Servicios."],
        ["b_asig_pct",     "%",        "0",      "input",          "% infra global asignado."],
        ["b_amort_anos",   "años",     "10",     "input",          "Vida útil infra."],
        ["b_mantenimiento","USD/año",  "0",      "input",          "Mantenimiento anual."],
        ["b_pc",           "USD/kg",   "1.40",   "input",          "Compra recría (si 1ª activa con compra externa)."],
        ["r_precio_venta", "USD/kg",   "2.30",   "input",          "Precio si Recría es salida."],
        ["r_flete_entrada","USD/cab",  "4",      "input",          "Flete de entrada."],
        ["r_flete_salida", "USD/cab",  "6",      "input",          "Flete de salida."],
        ["b_comision_pct", "%",        "2.5",    "input",          "Comisión."],
        ["r_otros",        "USD/cab",  "8",      "input",          "Otros gastos."],
    ],
    col_widths=[3.7, 2.0, 1.8, 2.4, 6.2],
)

add_h2(doc, "3.4 Engorde (C) — 220 → 430 kg, 180 días")
add_table(doc,
    headers=["Parámetro", "Unidad", "Default", "Tipo", "Descripción"],
    rows=[
        ["c_kg_entrada",   "kg",       "220",    "input/herencia", "Editable si 1ª activa; sino hereda."],
        ["c_dias",         "días",     "180",    "input",          "Duración."],
        ["t_peso_final",   "kg",       "430",    "input",          "Peso de terminación."],
        ["t_gdp",          "kg/día",   "1.167",  "input",          "GDP en feedlot."],
        ["t_ca",           "kg MS/kg", "7.5",    "input",          "Conversión."],
        ["t_mortalidad",   "%",        "2.0",    "input",          "Mortandad."],
        ["t_sanidad",      "USD/cab",  "30",     "input",          "Sanidad."],
        ["t_mo_mes",       "USD/mes",  "0",      "input",          "MO."],
        ["c_combustible",  "USD/mes",  "0",      "input",          "Combustible."],
        ["c_servicios",    "USD/mes",  "0",      "input",          "Servicios."],
        ["c_asig_pct",     "%",        "0",      "input",          "% infra asignado."],
        ["c_amort_anos",   "años",     "10",     "input",          "Vida útil."],
        ["c_mantenimiento","USD/año",  "0",      "input",          "Mantenimiento."],
        ["c_pc",           "USD/kg",   "2.30",   "input",          "Compra recriado (si 1ª activa)."],
        ["t_precio_venta", "USD/kg",   "3.50",   "input",          "Precio gancho mercado interno."],
        ["t_flete_entrada","USD/cab",  "5",      "input",          "Flete entrada."],
        ["t_flete_salida", "USD/cab",  "12",     "input",          "Flete salida."],
        ["c_comision_pct", "%",        "2.5",    "input",          "Comisión."],
        ["t_amortizacion", "USD/cab",  "18",     "input",          "Amortización legacy de corrales."],
        ["t_otros",        "USD/cab",  "10",     "input",          "Otros."],
    ],
    col_widths=[3.7, 2.0, 1.8, 2.4, 6.2],
)

add_h2(doc, "3.5 Estructura (global)")
add_table(doc,
    headers=["Parámetro", "Unidad", "Descripción"],
    rows=[
        ["infra_valor_total", "USD",   "Inversión total en infraestructura. Se asigna parcialmente vía *_asig_pct."],
        ["a_/b_/c_asig_pct",  "%",     "% del valor total imputado a cada etapa. La suma puede ser <100%."],
        ["a_/b_/c_amort_anos","años",  "Vida útil — divide el adjudicado para obtener amortización anual."],
        ["a_/b_/c_mantenimiento","USD/año","Mantenimiento anual absoluto, no porcentual."],
    ],
    col_widths=[5.0, 2.0, 9.0],
)

add_h2(doc, "3.6 Tabla de feed (bioeconómico)")
add_p(doc,
      "Cada etapa expone una tabla de ingredientes editable con dos "
      "columnas: % de inclusión y precio USD/kg MS. El motor calcula el "
      "precio ponderado de la ración como Σ(pct × precio) / Σ(pct). Es la "
      "única vía para configurar el costo de alimentación.")

add_pagebreak(doc)


# ── 4. FÓRMULAS ────────────────────────────────────────────────────────────

add_h1(doc, "4. Fórmulas")
add_p(doc,
      "Esta sección consolida todas las fórmulas del modelo con sus "
      "variables, unidades e interpretación.")

add_h2(doc, "4.1 Compra de hacienda")
add_formula(doc, "compra_cab = pc × kg_in")
add_p(doc,
      "pc = precio_compra global cuando Cría es 1ª activa, b_pc cuando "
      "Recría es 1ª activa, c_pc cuando Engorde es 1ª activa. "
      "Unidad: USD/cab. Sólo aplica a la 1ª etapa activa con compra externa.",
      italic=False, size=10)

add_h2(doc, "4.2 Alimentación bioeconómica")
add_formula(doc, "kg_carne     = max(kg_out − kg_in, 0)")
add_formula(doc, "consumo_MS   = kg_carne × CA                (kg MS/cab)")
add_formula(doc, "precio_pond  = Σ(pct_i × precio_i) / Σ(pct_i)   (USD/kg MS)")
add_formula(doc, "alim_cab     = consumo_MS × precio_pond        (USD/cab)")
add_p(doc,
      "kg_carne sólo cuenta la ganancia neta. CA es la conversión "
      "alimenticia (kg de materia seca por kg de peso vivo ganado). "
      "precio_pond pondera los ingredientes por su % de inclusión en la "
      "ración. La alimentación nunca se ingresa como USD/cab/día manual.",
      italic=False, size=10)

add_h2(doc, "4.3 Operación (USD/mes → USD/cab)")
add_formula(doc, "op_total_ciclo = (mo_mes + comb_mes + serv_mes) / 30 × días")
add_formula(doc, "op_cab         = op_total_ciclo / cabezas")
add_p(doc,
      "Los tres rubros se ingresan como USD/mes absolutos (no por cabeza). "
      "Se prorratean a USD/día y al ciclo, y luego se reparten entre el "
      "número de cabezas que la etapa absorbe.",
      italic=False, size=10)

add_h2(doc, "4.4 Estructura (amortización + mantenimiento)")
add_formula(doc, "adjudicado    = infra_valor_total × asig%/100")
add_formula(doc, "amort_anual   = adjudicado / amort_anos")
add_formula(doc, "amort_ciclo   = amort_anual × días / 365")
add_formula(doc, "mant_ciclo    = mantenimiento_anual × días / 365")
add_formula(doc, "estr_cab      = (amort_ciclo + mant_ciclo) / cabezas")
add_p(doc,
      "La infra es una inversión global (un único activo) que se asigna por "
      "porcentajes a cada etapa. El % no necesariamente suma 100 (puede "
      "haber actividades fuera del modelo). El mantenimiento es absoluto "
      "USD/año por etapa.",
      italic=False, size=10)

add_h2(doc, "4.5 Comercialización")
add_formula(doc, "comision_cab = (com%/100) × precio_venta × kg_out")
add_formula(doc, "com_cab      = comision_cab + flete_entrada + flete_salida")
add_p(doc,
      "Comisión sobre ingreso bruto unitario. Fletes son USD/cab fijos "
      "por etapa.",
      italic=False, size=10)

add_h2(doc, "4.6 Costo financiero")
add_formula(doc, "capital     = compra + alim + sanidad + op + estr + com")
add_formula(doc, "financiero  = capital × tasa%/100 × días/365")
add_p(doc,
      "Se aplica sobre el capital total inmovilizado durante el ciclo, no "
      "sólo sobre la compra. Días normalizados a una base anual.",
      italic=False, size=10)

add_h2(doc, "4.7 Mortandad")
add_formula(doc, "mortandad_cab = (mort%/100) × kg_out × precio_venta")
add_formula(doc, "cab_vend      = round( cab_in × (1 − mort/100) )")
add_p(doc,
      "Doble lectura: en la slide Costos, mortandad se contabiliza como "
      "‘ingreso perdido’ para visualizarla como bucket de costo de "
      "oportunidad. En Margen Bruto, NO se suma al costo — la asimetría "
      "entre cab_in (incurre el costo) y cab_vend (genera ingreso) ya "
      "captura el impacto, evitando doble contabilización.",
      italic=False, size=10)

add_h2(doc, "4.8 Ingresos")
add_formula(doc, "ingreso_cab   = kg_out × precio_venta")
add_formula(doc, "ingreso_total = ingreso_cab × cab_vend")

add_h2(doc, "4.9 Margen bruto")
add_formula(doc, "costo_cab          = capital + financiero")
add_formula(doc, "costo_total        = costo_cab × cab_in")
add_formula(doc, "margen_bruto_total = ingreso_total − costo_total")
add_formula(doc, "margen_bruto_cab   = margen_bruto_total / cab_in")
add_formula(doc, "kg_prod_total      = (kg_out − kg_in) × cab_vend")
add_formula(doc, "margen_kg          = margen_bruto_total / kg_prod_total")
add_formula(doc, "USD/cab/día        = margen_bruto_cab / días")

add_h2(doc, "4.10 Cabezas en cascada")
add_formula(doc, "cabezas[1ª activa] = n_terneros")
add_formula(doc, "cabezas[i+1]       = round( cabezas[i] × (1 − mort_i/100) )")
add_p(doc,
      "La cascada respeta sólo el slice activo: si Cría está OFF, Recría "
      "(1ª activa) parte de n_terneros. Etapas inactivas no consumen "
      "cabezas.",
      italic=False, size=10)

add_h2(doc, "4.11 Encadenamiento de pesos")
add_formula(doc, "kg_in[1ª activa] = K.*_KG_ENTRADA  (input editable)")
add_formula(doc, "kg_in[i+1]       = kg_out[i]        (heredado)")

add_pagebreak(doc)


# ── 5. SENSIBILIDAD Y ROBUSTEZ ─────────────────────────────────────────────

add_h1(doc, "5. Sensibilidad y robustez")

add_h2(doc, "5.1 Tornado")
add_p(doc,
      "Para cada etapa activa se evalúa el margen por cabeza con un "
      "delta positivo y otro negativo aplicado a una variable a la vez. "
      "El swing es la diferencia absoluta entre el resultado +Δ y −Δ. Las "
      "barras se ordenan de menor a mayor swing.")
add_table(doc,
    headers=["Variable", "Delta", "Cálculo del override"],
    rows=[
        ["Precio compra", "±20 %",  "compra_cab ×= (1 ± 0.20)"],
        ["Precio venta",  "±20 %",  "pv ×= (1 ± 0.20); afecta ingreso y comisión"],
        ["Precio maíz",   "±20 %",  "alim ×= (1 ± 0.20)"],
        ["GDP",           "±15 %",  "kg_carne ×= (1 ± 0.15); alim escala con kg_carne"],
        ["Mortandad",     "±5 pp",  "mort_pct += ±5; cap [0, 99]"],
        ["Flete",         "±20 %",  "(flete_entrada + flete_salida) ×= (1 ± 0.20)"],
    ],
    col_widths=[3.5, 2.5, 10.0],
)

add_h2(doc, "5.2 Breakeven analítico")
add_p(doc,
      "Para cada etapa se calculan los puntos de equilibrio (margen = 0):")
add_formula(doc, "precio_equilibrio = costo_total / kg_vendidos")
add_formula(doc, "mort_max_pct      = 100 × (1 − base / ingreso_cab)")
add_formula(doc, "α (alim factor)   = ((1−m) × ingreso_cab − (base − alim)) / alim")
add_formula(doc, "precio_alim_max   = α × precio_alim_actual")
add_formula(doc, "ca_max            = α × ca_actual")
add_formula(doc, "gdp_min           = (X_min − kg_in) / días,  X_min = K_const / (pv × ((1−m) − com_share))")
add_p(doc,
      "Donde base = capital + financiero (sin sumar mortandad). El "
      "breakeven analítico asume que el costo financiero no se recalcula "
      "al variar kg_out (aproximación lineal).",
      italic=False, size=10)

add_h2(doc, "5.3 Score de robustez")
add_p(doc,
      "Robustez = 100 − composite, con composite ponderada de cinco "
      "factores de riesgo (mayor factor = más riesgoso):")
add_table(doc,
    headers=["Factor", "Cálculo", "Cap"],
    rows=[
        ["Sensibilidad maíz",  "swing(maíz) / |baseline| × 50",   "100"],
        ["Volatilidad",        "Σ swings / |baseline| × 30",       "100"],
        ["Duración",           "días / 7.30",                      "100"],
        ["Capital relativo",   "costo_total / max(activos) × 100", "100"],
        ["Mortandad",          "mort% × 10",                       "100"],
    ],
    col_widths=[4.5, 7.0, 1.5],
)
add_p(doc,
      "Los pesos del composite están definidos en _RISK_WEIGHTS. El score "
      "robustez es el complemento: 100 puntos cuando el riesgo composite "
      "es 0, 0 puntos cuando el riesgo composite alcanza 100.",
      italic=False, size=10)

add_h2(doc, "5.4 Alertas")
add_p(doc,
      "Las alertas se generan focalizadas en dos condiciones: (a) etapas "
      "con robustez baja, y (b) variables del tornado con swing alto "
      "respecto del baseline. Cada alerta identifica la variable + la "
      "etapa, evitando reportar redundancias entre slides.")

add_pagebreak(doc)


# ── 6. REPORTES ────────────────────────────────────────────────────────────

add_h1(doc, "6. Reportes — One pager A4")

add_h2(doc, "6.1 Estructura visual")
add_bullets(doc, [
    "Toolbar editable: empresa, responsable, fecha, toggle de logo.",
    "Header A4: logo (izquierda) + datos del reporte (derecha).",
    "Título centrado: ‘Reporte Estratégico Ganadero’.",
    "Hero del resumen ejecutivo: indica el sistema productivo integrado y "
    "muestra el margen total como headline.",
    "4 KPI cards alineadas: Margen sistema, USD/cab/día, Riesgo (promedio), "
    "Robustez (promedio).",
    "Fila 1 de mini-bloques: Modelo Productivo, Costos, Ingresos.",
    "Fila 2 de mini-bloques: Margen Bruto, Sensibilidad y Riesgo.",
])

add_h2(doc, "6.2 Indicadores agregados del sistema")
add_table(doc,
    headers=["Indicador", "Cálculo"],
    rows=[
        ["Margen total",         "Σ margen_bruto_total de las 3 etapas"],
        ["USD / cab / día",      "promedio simple de las 3 etapas"],
        ["Riesgo (promedio)",    "promedio simple de risk_composite por etapa"],
        ["Robustez (promedio)",  "promedio simple de robustness por etapa"],
    ],
    col_widths=[4.5, 11.5],
)
add_p(doc,
      "El reporte agrega métricas del sistema entero. Para foco por etapa, "
      "los mini-bloques exponen el resumen de cada slide.",
      italic=False, size=10)

add_h2(doc, "6.3 Mini-bloques (por slide)")
add_table(doc,
    headers=["Mini-bloque", "Bullets principales", "Mini gráfico"],
    rows=[
        ["Modelo productivo",   "Días totales · peso final · GDP promedio",            "Sparkline kg vs días"],
        ["Costos",              "Top etapa por costo · USD/kg · % por categoría",      "Hbars top 4 categorías"],
        ["Ingresos",            "Ingreso sistema · top etapa · cabezas vendidas",      "Hbars ingresos por etapa"],
        ["Margen bruto",        "Top etapa por margen/cab · margen/kg · margen total", "Hbars margen por etapa"],
        ["Sensibilidad y riesgo","Variable crítica · precio equilibrio · robustez+riesgo","Hbars top swings tornado"],
    ],
    col_widths=[3.8, 7.5, 4.7],
)


# ── Cierre ─────────────────────────────────────────────────────────────────

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("— Fin del documento —")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = C_SUBTLE


# ── Guardar ────────────────────────────────────────────────────────────────
doc.save(str(OUT_PATH))
print(f"OK: {OUT_PATH}")
